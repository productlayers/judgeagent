#!/usr/bin/env python3
"""Generate Agent-as-Judge team doc (Word)."""

from __future__ import annotations

import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for j, h in enumerate(headers):
        table.rows[0].cells[j].text = h
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            table.rows[i].cells[j].text = val


def main() -> None:
    doc = Document()

    title = doc.add_heading("Agent-as-Judge: Current vs Proposed", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.add_run(
        "For teammates — what's on GitHub today, what we're building toward, "
        "and how the pieces fit.\n"
    ).italic = True
    doc.add_paragraph("Repo: https://github.com/productlayers/judgeagent")
    doc.add_paragraph("Latest remote commit: fe55e3a — Add CopilotKit LLM judge UI")
    doc.add_paragraph("If local is behind: run git pull")

    doc.add_heading("One-sentence summary", level=1)
    doc.add_paragraph(
        "Today: A Python handoff-verification judge (MCP + CLI) with Weave tracing, "
        "plus a separate CopilotKit UI for generic LLM response scoring — not connected."
    )
    doc.add_paragraph(
        "Proposed: A LangGraph multi-agent orchestration demo where handoffs are verified "
        "by the Python judge, state flows through Redis, the CopilotKit UI shows gap analysis "
        "and recommendations, and Weave traces the full run."
    )

    doc.add_page_break()
    doc.add_heading("Part 1: What's on GitHub Today", level=1)

    doc.add_heading("Backend (Python) — Handoff verification", level=2)
    doc.add_paragraph(
        'Purpose: Verify whether an agent\'s "I\'m done" claim is supported by '
        "evidence at a multi-agent handoff."
    )
    doc.add_paragraph("Pipeline:")
    for step in [
        "derive_rubric — LLM turns goal into criteria",
        "evaluate_rubric — LLM checks criteria vs evidence",
        "aggregate_verdict — code returns pass / fail / uncertain",
    ]:
        doc.add_paragraph(step, style="List Bullet")

    add_table(
        doc,
        ["Surface", "What it does"],
        [
            ("MCP server (judge-mcp)", "judge_handoff tool over stdio"),
            ("CLI (judge-cli)", "JSON in → verdict out; exit codes for CI"),
            ("Python API", "run_judge(JudgeInput)"),
        ],
    )

    doc.add_paragraph(
        "Input (JudgeInput): original_goal, agent_claim, evidence, "
        "agent_system_prompt, workflow_context"
    )
    doc.add_paragraph(
        "Output (JudgeResult): verdict, can_continue, rubric, evaluations, "
        "failure_reasons, prompt_improvements"
    )
    doc.add_paragraph("Demos: travel_demo.py (fail → pass); run_evals.py (3 guard cases)")
    doc.add_paragraph(
        "Not built yet: LangGraph, Redis, UI connection, Weave Evals, trace-ID pull"
    )

    doc.add_heading("Frontend (CopilotKit) — LLM judge workbench", level=2)
    doc.add_paragraph(
        "Purpose: Score and compare candidate LLM responses vs prompt + reference + rubric."
    )
    doc.add_paragraph("Stack: Next.js + CopilotKit v2 in app/")
    code = doc.add_paragraph()
    code.add_run(
        "npm install\ncp .env.example .env.local\nnpm run dev   # http://localhost:3000"
    ).font.name = "Courier New"
    doc.add_paragraph(
        "Features: Editable rubric, candidates, score breakdown, CopilotKit sidebar."
    )
    p = doc.add_paragraph()
    p.add_run("Important: ").bold = True
    p.add_run("Scoring is client-side heuristics — does NOT call Python run_judge().")
    doc.add_paragraph(
        "Not built yet: Flow canvas, handoff gap analysis, LangGraph, Redis, Weave"
    )

    doc.add_heading("Weave — Already in Python Backend", level=2)
    p = doc.add_paragraph()
    p.add_run("Weave is in the Python judge today. It is NOT in the CopilotKit UI.").bold = True
    code2 = doc.add_paragraph()
    code2.add_run(
        "import weave\n"
        'weave.init("agent-judge")   # tracks LLM calls\n\n'
        "@weave.op()\n"
        "def derive_rubric(data): ..."
    ).font.name = "Courier New"
    doc.add_paragraph("Enable: WEAVE_PROJECT=agent-judge + WANDB_API_KEY in .env")
    doc.add_paragraph("Quickstart: wandb.me/use-weave")
    trace = doc.add_paragraph()
    trace.add_run(
        "run_judge → derive_rubric → evaluate_rubric → aggregate_verdict "
        "(+ nested LLM calls)"
    ).font.name = "Courier New"

    doc.add_page_break()
    doc.add_heading("Part 2: Proposed Structure", level=1)
    doc.add_paragraph(
        "Product: Agent Flow Audit — LangGraph flow + judge at handoffs + "
        "CopilotKit audit UI + Redis + Weave."
    )

    doc.add_heading("Proposed Backend", level=2)
    for item in [
        "LangGraph orchestrates Planner → Judge gate → Booker",
        "Python judge as graph node at handoff edges",
        "New schemas: HandoffAuditRequest, LangGraphHandoffContext, EvidenceBundle, WeaveRef",
        "Historical audit — verify what ran, don't re-run agent by default",
        "Redis: run state, pub/sub, pass-rate stats, HITL gates",
        "Weave: trace LangGraph nodes + judge; Evals for regression",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Proposed Frontend", level=2)
    for item in [
        "Flow canvas — click handoff to audit",
        "Gap analysis — intended (rubric) vs actual (evaluations)",
        "Recommendations with prompt patches + re-audit",
        "Pass rates from Redis; HITL for uncertain verdicts",
        "Weave trace links; call Python run_judge() instead of heuristics",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Proposed Weave Integration", level=2)
    add_table(
        doc,
        ["Layer", "Today", "Proposed"],
        [
            ("Judge pipeline", "@weave.op on all stages", "Same"),
            ("LLM calls", "Auto-traced via weave.init()", "Same"),
            ("LangGraph nodes", "Not yet", "@weave.op on all nodes"),
            ("Historical data", "Manual evidence string", "Optional pull from trace_id"),
            ("Evals", "Not yet", "Weave Evals on guard cases"),
            ("UI", "Not yet", "Trace links in audit panel"),
        ],
    )

    doc.add_page_break()
    doc.add_heading("Part 3: Diff Summary", level=1)

    doc.add_heading("Backend", level=2)
    add_table(
        doc,
        ["Area", "GitHub Today", "Proposed"],
        [
            ("Core judge", "Python MCP/CLI", "Keep + LangGraph adapter"),
            ("Orchestration", "Script demo", "LangGraph with judge gate"),
            ("Schemas", "JudgeInput / JudgeResult", "+ HandoffAuditRequest, etc."),
            ("Historical runs", "Manual JSON", "Graph state + Weave trace pull"),
            ("Redis", "No", "Run state, pub/sub, stats"),
            ("Weave", "Judge only", "Judge + LangGraph + Evals"),
            ("Re-run agent", "No", "Optional only"),
        ],
    )

    doc.add_heading("Frontend", level=2)
    add_table(
        doc,
        ["Area", "GitHub Today", "Proposed"],
        [
            ("App", "LLM Judge workbench", "Agent Flow Audit cockpit"),
            ("Use case", "Score LLM answers", "Audit agent handoffs"),
            ("Scoring", "Client heuristics", "Python run_judge()"),
            ("Rubric", "Manual in UI", "Auto-derived from goal"),
            ("Flow viz", "No", "LangGraph canvas"),
            ("Gap analysis", "Heuristic strengths/risks", "JudgeResult evaluations"),
            ("Live updates", "No", "Redis pub/sub"),
            ("Weave", "No", "Trace links"),
            ("Python connected", "No", "Yes"),
        ],
    )

    doc.add_heading("The Main Gap Today", level=2)
    doc.add_paragraph(
        "Two separate products in one repo: Python handoff judge (Weave yes) and "
        "CopilotKit LLM workbench (Weave no), not wired together. Proposed: LangGraph "
        "orchestrates, Python judges, CopilotKit displays, Redis coordinates, Weave records."
    )

    doc.add_heading("Part 4: Suggested Team Split", level=1)
    add_table(
        doc,
        ["Track", "Deliverables"],
        [
            ("Backend / orchestration", "LangGraph demo, judge gate node, Redis, schemas"),
            ("Frontend", "Rewire CopilotKit to gap analysis, flow canvas, HITL"),
            ("Judge core", "Keep stable in src/agent_judge/"),
            ("Weave / evals", "Trace all nodes; eval dataset"),
            ("Sponsor story", "Redis + CopilotKit + Weave demo"),
        ],
    )

    doc.add_heading("Part 5: Build Order", level=1)
    for i, step in enumerate(
        [
            "LangGraph travel pipeline",
            "Wire judge gate to run_judge()",
            "Weave on LangGraph nodes",
            "CopilotKit shows JudgeResult gap table",
            "Redis live updates",
            "Weave Evals",
        ],
        start=1,
    ):
        doc.add_paragraph(f"{i}. {step}", style="List Number")

    doc.add_heading("Quick Reference", level=1)
    add_table(
        doc,
        ["Question", "Answer"],
        [
            ("Weave in repo?", "Yes — Python judge; set WEAVE_PROJECT"),
            ("Weave in UI?", "No"),
            ("LangGraph?", "Proposed"),
            ("Redis?", "Proposed"),
            ("UI connected to Python?", "No — proposed"),
            ("Teammates should run?", "git pull (fe55e3a)"),
        ],
    )

    root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    out = os.path.join(root, "docs", "Agent-as-Judge-Current-vs-Proposed.docx")
    os.makedirs(os.path.dirname(out), exist_ok=True)
    doc.save(out)
    print(f"Created: {out}")


if __name__ == "__main__":
    main()
